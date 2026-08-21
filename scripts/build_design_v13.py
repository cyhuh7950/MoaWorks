from copy import deepcopy
from pathlib import Path

from docx import Document


SOURCE = Path(r"D:\Project\MoaWorks\docs\MoaWorks_Groupware_Design_v1.2.docx")
TARGET = Path(r"D:\Project\MoaWorks\docs\MoaWorks_Groupware_Design_v1.3.docx")


SECTION_22 = [
    ("Heading 1", "22. 단계별 진행계획"),
    ("Heading 2", "22.0 전제"),
    ("Normal", "MoaWorks는 server, admin-web, user-web, desktop-client, mobile-app의 5개 프로그램 단위로 개발한다."),
    ("Normal", "server는 공통 인증, 권한, 조직, 메일, 전자결재, 알림, 감사 로그, 운영 감시, 번역, 외부 연동 API를 담당한다."),
    ("Normal", "admin-web은 운영자 전용 화면이며 초기 설정, 사용자·조직·도메인·발신 엔진 관리, 상태 점검, 감사 로그 확인, 운영 정책 제어를 담당한다."),
    ("Normal", "user-web, desktop-client, mobile-app은 일반 사용자용 업무 클라이언트이며 모두 동일한 서버 API만 호출한다."),
    ("Normal", "모든 운영 데이터의 최종 저장소는 PostgreSQL이다."),
    ("Normal", "모든 개발 확인은 WSL 실행 환경에 실제로 서비스를 띄운 상태에서 진행한다. 코드 검토만으로 완료 판정하지 않는다."),
    ("Normal", "단계 중 임시 검증을 위해 파일 기반 목업 또는 메모리 저장을 사용할 수 있으나, 단계 3 종료 전까지 반드시 제거되어야 하며 운영 산출물로 인정하지 않는다."),
    ("Normal", "단계 완료 판정은 문서, 코드, WSL 실행 검증 3가지를 모두 충족해야 한다."),
    ("Normal", "개발 진행 순서는 작업 지시, 작업자 보고, 설계·테스트 확인 담당자의 WSL 실행 검증, 설계·테스트 확인 보고, 신산님 확인, 재작업 또는 다음 단계 진행 순서로 고정한다."),
    ("Heading 2", "22.1 단계 0. 기준 확정 및 구현 금지선 정의"),
    ("Normal", "목적: 제품 범위, 저장 구조, 프로그램 역할, 운영 원칙, 완료 판정 기준을 먼저 고정하여 이후 개발자가 임의 해석으로 구조를 바꾸지 못하게 한다."),
    ("Normal", "작업 범위: MVP/P1/P2 범위, 5개 프로그램 역할, 핵심 테이블, 핵심 API, 권한 경계, 감사 로그, 운영 감시, 배포 및 승인 원칙, WSL 검증 원칙 확정."),
    ("Normal", "작업 내용: PostgreSQL을 유일한 운영 DB로 고정하고, 파일 기반 저장은 운영 구조로 인정하지 않는다는 금지선을 문서로 명시한다. 핵심 테이블 정의서, 핵심 API 정의서, 권한 정책서, 운영 이벤트 정의서, 릴리즈 승인 기준서, WSL 검증 절차를 작성한다."),
    ("Normal", "산출물: 제품 범위표, 시스템 컨텍스트 문서, 핵심 테이블 정의서, 핵심 API 정의서, 운영 이벤트 정의서, 권한 정책서, 승인 기준서, WSL 검증 절차서."),
    ("Normal", "완료 기준: companies, users, departments, roles, mail_accounts, mail_provider_configs, approval_documents, approval_lines, notifications, audit_logs, monitoring_events가 핵심 운영 테이블로 문서에 고정되어 있어야 한다. PostgreSQL 운영 원칙, 파일 기반 임시구현 종료 기준, WSL 검증 순서가 명시되어야 한다."),
    ("Normal", "검수 기준: 구현자가 추가 해석 없이 저장 구조, 프로그램 경계, 검증 순서를 이해할 수 있어야 한다."),
    ("Normal", "작업 지시 단위: 제품 범위 담당, 테이블 정의 담당, API 계약 담당, 운영 기준 담당, 승인 기준 담당."),
    ("Normal", "차단 조건: DB 없이도 완료로 판정할 수 있게 문서가 열려 있으면 단계 통과 불가."),
    ("Heading 2", "22.2 단계 1. 실행 골격과 설치 흐름 구축"),
    ("Normal", "목적: 운영자가 브라우저에서 설치를 시작하고 서버 상태를 확인할 수 있는 최소 실행 구조를 만든다."),
    ("Normal", "작업 범위: server, admin-web, user-web, desktop-client, mobile-app 실행 골격, 초기 설정 흐름, Health API, 배포 골격."),
    ("Normal", "작업 내용: FastAPI 기반 server 골격, 관리자 웹 초기 설정 화면, 일반 사용자 3개 클라이언트 골격, Docker Compose 초안, /health, /setup/validate, /setup/initialize API 골격을 구현한다."),
    ("Normal", "산출물: 5개 프로그램 골격, 초기 설정 Wizard 1차, Compose 초안, 실행 가이드 초안, WSL 기동 로그."),
    ("Normal", "완료 기준: server와 admin-web이 WSL에서 실제 기동되고, 관리자 웹에서 초기 설정 화면 접근이 가능해야 한다. 일반 사용자용 3개 클라이언트의 디렉터리와 앱 골격이 고정되어 있어야 한다."),
    ("Normal", "검수 기준: 작업자 보고 후 설계·테스트 확인 담당자가 WSL에서 동일 명령으로 재기동 및 접속 확인을 끝내고 보고해야 하며, 설치 실패 시 반쯤 설치된 상태로 남지 않아야 한다."),
    ("Normal", "작업 지시 단위: server 골격 담당, admin-web 골격 담당, user-web 골격 담당, desktop-client 골격 담당, mobile-app 골격 담당, deploy 골격 담당."),
    ("Normal", "차단 조건: 이후 DB 전환을 고려하지 않은 임시 라우팅 또는 임시 데이터 구조면 통과 불가."),
    ("Heading 2", "22.3 단계 2. PostgreSQL 스키마와 영속 계층 구축"),
    ("Normal", "목적: 임시 저장이 아니라 실제 운영 데이터를 PostgreSQL에 저장하는 구조를 먼저 완성한다."),
    ("Normal", "작업 범위: DB 스키마, ORM 또는 SQL 계층, 마이그레이션, 초기 설정 저장, 사용자·조직·권한 저장 구조."),
    ("Normal", "작업 내용: PostgreSQL 15+ 기준 스키마와 마이그레이션을 구현하고, setup/initialize 결과를 DB에 저장한다. 사용자, 부서, 권한, 메일 계정, Relay 설정을 DB에 저장하도록 전환한다. /health에서 실제 DB 연결을 검사한다."),
    ("Normal", "산출물: DDL 또는 마이그레이션 세트, DB 접근 계층, DB 초기화 절차, DB 연결 헬스체크, WSL DB 검증 로그."),
    ("Normal", "완료 기준: 회사, 관리자, 기본 조직, 기본 권한, 기본 메일 설정이 PostgreSQL에 실제 저장되고 WSL에서 서버 재시작 후에도 유지되어야 한다. 파일 기반 상태 저장은 운영 경로에서 제거되어야 한다."),
    ("Normal", "검수 기준: 작업자 보고 후 설계·테스트 확인 담당자가 WSL에서 저장, 조회, 재기동 유지, 마이그레이션 재실행, DB 연결 실패 시 경고 동작을 직접 검증해야 한다."),
    ("Normal", "작업 지시 단위: DB 스키마 담당, 마이그레이션 담당, 저장 계층 담당, 헬스체크 담당, 운영 문서 담당."),
    ("Normal", "차단 조건: API는 동작하지만 DB에 저장되지 않거나 서버 재시작 후 상태가 유실되면 통과 불가."),
    ("Heading 2", "22.4 단계 3. 계정·조직·권한·인증 운영 기반 완성"),
    ("Normal", "목적: 운영자가 사람과 조직을 관리하고, 모든 클라이언트가 동일 인증 체계로 로그인하도록 만든다."),
    ("Normal", "작업 범위: 사용자·조직·권한 CRUD, 메일 계정 자동 생성, 도메인 검증, Relay 테스트, 공통 로그인, 권한 재검증."),
    ("Normal", "작업 내용: 관리자용 사용자·부서·권한 관리 API와 화면을 구현하고, 사용자 생성 시 메일 계정이 자동 생성되도록 한다. 공통 로그인, 토큰 발급, 요청 시점 권한 재검증을 구현한다."),
    ("Normal", "산출물: 계정·조직·권한 API, 관리자 사용자 관리 화면, 공통 인증 계약서, 도메인 검증 기능, Relay 테스트 기능, WSL 로그인 검증 로그."),
    ("Normal", "완료 기준: 관리자 웹에서 사용자 생성, 수정, 비활성화가 가능해야 하고, 비활성 사용자 또는 비활성 권한은 로그인 후에도 즉시 차단되어야 한다. 세 클라이언트가 동일 로그인 API 명세를 사용해야 한다."),
    ("Normal", "검수 기준: 작업자 보고 후 설계·테스트 확인 담당자가 WSL에서 관리자 로그인, 일반 사용자 로그인, 비활성 사용자 차단, 권한 재검증을 직접 확인해야 한다."),
    ("Normal", "작업 지시 단위: 계정 API 담당, 조직 API 담당, 인증 담당, 관리자 화면 담당, 도메인 및 Relay 담당."),
    ("Normal", "차단 조건: 토큰 내부 권한만 믿고 실시간 상태 재검증이 없으면 통과 불가."),
    ("Heading 2", "22.5 단계 4. 전자결재 핵심 업무 완성"),
    ("Normal", "목적: 전자결재를 실제 업무에 사용할 수 있는 수준으로 완성한다."),
    ("Normal", "작업 범위: 결재 문서, 결재선, 상태 머신, 감사 로그, 관리자 직권 처리, 사용자 결재 UI."),
    ("Normal", "작업 내용: approval_documents, approval_lines, audit_logs를 DB에 구현하고, 작성, 상신, 승인, 반려, 회수, 재기안 API를 구현한다. 관리자 직권 승인 또는 반려 API를 별도 권한 스코프로 분리한다. user-web과 desktop-client, mobile-app에 동일 API 기반 결재 화면을 붙인다."),
    ("Normal", "산출물: 전자결재 API, 상태 전이 규칙 문서, 감사 로그 조회 기능, 3클라이언트 결재 UI 1차, WSL 결재 검증 로그."),
    ("Normal", "완료 기준: 허용된 상태 전이만 가능해야 하고, 승인 완료 문서는 재처리되지 않아야 하며, 모든 상태 변경이 감사 로그에 남아야 한다. 세 클라이언트는 동일 결재 API로 조회와 처리가 가능해야 한다."),
    ("Normal", "검수 기준: 작업자 보고 후 설계·테스트 확인 담당자가 WSL에서 작성, 상신, 승인, 반려, 회수, 재기안, 직권 처리, 동시성 충돌 통제를 직접 검증해야 한다."),
    ("Normal", "작업 지시 단위: 결재 스키마 담당, 상태 머신 담당, 결재 API 담당, user-web 결재 UI 담당, desktop-client 결재 UI 담당, mobile-app 결재 UI 담당, 감사 로그 담당."),
    ("Normal", "차단 조건: 결재 결과가 DB와 감사 로그 사이에서 불일치하면 통과 불가."),
    ("Heading 2", "22.6 단계 5. 메일·알림·운영 감시 통합"),
    ("Normal", "목적: 메일 이벤트, 결재 이벤트, 시스템 이상 징후를 역할에 맞게 저장하고 전달한다."),
    ("Normal", "작업 범위: 알림 저장, 읽음 처리, 실시간 전달, Watcher, 관리자 대시보드, 클라이언트 알림 정책."),
    ("Normal", "작업 내용: notifications, monitoring_events 저장 구조를 구현하고, 결재, 메일, 시스템 이벤트 생성기를 만든다. SSE 또는 WebSocket 기반 전달 채널과 폴링 폴백을 구현한다. admin-web 운영 대시보드와 user-web, desktop-client, mobile-app 알림 UI를 구현한다."),
    ("Normal", "산출물: 알림 API, 운영 감시 API, 관리자 대시보드, 3클라이언트 알림 흐름, 알림 정책 문서, WSL 알림 검증 로그."),
    ("Normal", "완료 기준: 핵심 업무 이벤트가 사용자 클라이언트에 전달되고, 핵심 장애 이벤트가 관리자 웹에 표시되어야 한다. 읽음, 미확인, 재조회 흐름이 동작해야 한다."),
    ("Normal", "검수 기준: 작업자 보고 후 설계·테스트 확인 담당자가 WSL에서 관리자 경고, 사용자 업무 알림, 실시간 전달, 폴백, 재조회 흐름을 직접 검증해야 한다."),
    ("Normal", "작업 지시 단위: 알림 이벤트 담당, 전달 채널 담당, 운영 대시보드 담당, Watcher 담당, user-web 알림 UI 담당, desktop-client 알림 담당, mobile-app 알림 담당."),
    ("Normal", "차단 조건: 알림 누락, 중복, 권한 오전달이 있으면 통과 불가."),
    ("Heading 2", "22.7 단계 6. 다국어·번역·보안·백업/복구 완성"),
    ("Normal", "목적: 기본 기능을 해치지 않으면서 다국어와 번역을 얹고, 운영 안전성을 확보한다."),
    ("Normal", "작업 범위: 7개 언어 UI, 번역 Provider, 번역 장애 정책, 비밀정보 보호, 백업 및 복구 기준."),
    ("Normal", "작업 내용: admin-web과 user-web에 다국어를 적용하고 desktop-client, mobile-app의 locale 정책을 정렬한다. 번역 Provider 연결과 미연결 정책을 구현한다. 비밀번호와 비밀값을 암호화 저장하고, PostgreSQL 백업과 첨부파일 복구 절차를 정리한다."),
    ("Normal", "산출물: 다국어 UI 적용본, 번역 서비스, 보안 점검 결과서, 백업 및 복구 기준서, WSL 복구 검증 로그."),
    ("Normal", "완료 기준: 번역 기능이 없어도 핵심 그룹웨어 기능은 정상 동작해야 하고, 번역 장애가 로그인, 결재, 알림을 막지 않아야 한다. 비밀번호와 비밀값은 평문 저장되면 안 된다. DB 백업과 복구 절차가 실제 실행 가능한 문서로 존재해야 한다."),
    ("Normal", "검수 기준: 작업자 보고 후 설계·테스트 확인 담당자가 WSL에서 다국어 화면, 번역 미연결 정책, 비밀값 저장, 백업, 복구 절차를 직접 검증해야 한다."),
    ("Normal", "작업 지시 단위: 다국어 담당, 번역 서비스 담당, 보안 점검 담당, 백업 담당, 복구 담당, 운영 문서 담당."),
    ("Normal", "차단 조건: 번역 장애가 핵심 기능을 막거나 복구 문서가 실행 절차 없이 선언만 있으면 통과 불가."),
    ("Heading 2", "22.8 단계 7. 실운영형 통합 테스트와 성능 검증"),
    ("Normal", "목적: 기능 확인 수준이 아니라 운영 가능한 상태인지 검증한다."),
    ("Normal", "작업 범위: PostgreSQL 실연동, 메일·결재·알림 E2E, 장애 복구, 동시성, 100명 이하 규모 성능 검증."),
    ("Normal", "작업 내용: 실DB 기반 통합 테스트, 서버 재시작, DB 재시작, 네트워크 장애, 동시 로그인, 동시 결재, 알림 폭주, 느린 쿼리, 락 경합 시나리오를 수행한다. 세 클라이언트의 동일 API 정합성을 검증한다."),
    ("Normal", "산출물: 통합 테스트 시나리오서, 실행 로그, 성능 점검 보고서, 병목 및 보완 작업지시서, WSL 통합 검증 로그."),
    ("Normal", "완료 기준: PostgreSQL 기반 저장과 조회가 전 기능에서 확인되어야 하고, 장애 후 복구 절차가 실제 수행되어야 하며, 100명 이하 사용 기준의 응답성과 안정성 목표를 충족해야 한다."),
    ("Normal", "검수 기준: 작업자 보고 후 설계·테스트 확인 담당자가 WSL에서 같은 시나리오를 재실행해 결과를 확인하고 신산님에게 보고해야 한다. 파일기반 우회 테스트만 통과한 경우 인정하지 않는다."),
    ("Normal", "작업 지시 단위: 통합 테스트 담당, 성능 테스트 담당, 장애 복구 담당, 정합성 검증 담당, 병목 분석 담당."),
    ("Normal", "차단 조건: 실DB 테스트 미완료, 장애 복구 미검증, 성능 병목 미분석 상태면 통과 불가."),
    ("Heading 2", "22.9 단계 8. 배포 승인과 공개 운영 문서 완료"),
    ("Normal", "목적: 외부 사용자가 설치하고 운영할 수 있는 상태로 마감한다."),
    ("Normal", "작업 범위: 배포 스크립트, 설치 문서, 운영 문서, 관리자 가이드, 일반 사용자 가이드, 릴리즈 승인."),
    ("Normal", "작업 내용: 운영 환경 기준 Compose 또는 배포 절차를 정리하고, 관리자용 운영 문서와 일반 사용자용 웹, 설치형, 앱 사용 문서를 작성한다. 설치, 초기 설정, 로그인, 기본 업무 확인 절차를 문서와 실제 실행 결과로 맞춘다."),
    ("Normal", "산출물: 운영 배포 가이드, 설치 가이드, 관리자 사용 가이드, 일반 사용자 사용 가이드, 릴리즈 승인 보고서, WSL 문서 검증 로그."),
    ("Normal", "완료 기준: 문서만 보고 설치, 초기 설정, 관리자 로그인, 일반 사용자 로그인, 기본 업무 확인이 가능해야 한다. 배포 문서와 실제 실행 절차가 일치해야 한다."),
    ("Normal", "검수 기준: 설계·테스트 확인 담당자가 WSL에서 문서만 보고 직접 따라 실행한 뒤 결과를 신산님에게 보고해야 하며, 신산님 확인 전에는 완료로 판정하지 않는다."),
    ("Normal", "작업 지시 단위: 배포 담당, 설치 문서 담당, 관리자 문서 담당, 사용자 문서 담당, 릴리즈 승인 담당."),
    ("Normal", "차단 조건: 문서와 실제 동작이 다르거나 실운영 테스트 미완료 상태면 최종 완료 불가."),
    ("Heading 2", "22.10 최종 완료 판정"),
    ("Normal", "최종 완료는 다음 5가지를 모두 만족할 때만 인정한다."),
    ("List Bullet", "PostgreSQL 기반 영속 구조가 실제 동작한다."),
    ("List Bullet", "5개 프로그램의 역할과 API 경계가 설계서와 일치한다."),
    ("List Bullet", "관리자 기능과 일반 사용자 기능이 화면과 API에서 분리되어 있다."),
    ("List Bullet", "실운영형 통합 테스트와 복구 테스트가 완료되었다."),
    ("List Bullet", "배포, 설치, 운영 문서가 실제 실행 결과와 일치한다."),
    ("Heading 2", "22.11 단계별 작업 지시 원칙"),
    ("Normal", "각 작업지시는 반드시 아래 형식으로 작성한다."),
    ("List Bullet", "작업명"),
    ("List Bullet", "대상 프로그램"),
    ("List Bullet", "목적"),
    ("List Bullet", "입력 문서"),
    ("List Bullet", "선행 조건"),
    ("List Bullet", "구현 범위"),
    ("List Bullet", "제외 범위"),
    ("List Bullet", "저장 대상 테이블 또는 API"),
    ("List Bullet", "화면 또는 운영 흐름"),
    ("List Bullet", "산출물"),
    ("List Bullet", "완료 기준"),
    ("List Bullet", "검수 기준"),
    ("List Bullet", "실패 시 보완 지시"),
    ("List Bullet", "차단 조건"),
    ("Heading 2", "22.12 단계 진행 및 승인 순서"),
    ("Normal", "모든 단계는 아래 순서를 벗어나지 않는다."),
    ("List Number", "작업 지시: 신산님 또는 설계 기준 문서에 따라 작업 범위와 완료 기준을 고정한다."),
    ("List Number", "작업 보고: 작업 담당자가 구현 내용, 변경 파일, 실행 결과, 남은 의심 사항을 보고한다."),
    ("List Number", "설계·테스트 확인: 설계 및 테스트 확인 담당자가 WSL 실행 환경에 서비스를 직접 띄우고, 코드와 화면과 API를 실제로 확인한다."),
    ("List Number", "확인 보고: 설계·테스트 확인 담당자가 직접 확인한 사실만 근거로 통과, 보완 필요, 차단 상태를 보고한다."),
    ("List Number", "신산님 확인: 신산님이 보고 내용을 기준으로 재작업 또는 다음 단계 진행을 결정한다."),
    ("Normal", "설계·테스트 확인 담당자가 직접 실행하지 않은 내용은 확인 결과로 보고할 수 없다."),
]


def remove_paragraph(paragraph):
    element = paragraph._element
    parent = element.getparent()
    parent.remove(element)


def main():
    doc = Document(str(SOURCE))

    start_idx = None
    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if text == "v1.2":
            paragraph.text = "v1.3"
        elif text == "2026-06-18":
            paragraph.text = "2026-06-21"
        elif text.startswith("v1.2 주요 변경"):
            paragraph.text = "v1.3 주요 변경"
        elif "운영자용 관리자 웹과 일반 사용자용 웹·설치형·앱 구조 확정" in text:
            paragraph.text = "PostgreSQL 운영 저장 구조를 최종 기준으로 고정하고, 단계별 진행계획을 실운영 완료·통합테스트·배포 승인까지 확장"
        elif text == "22. 단계별 진행계획":
            start_idx = idx
            break

    if start_idx is None:
        raise RuntimeError("Section 22 start not found")

    for paragraph in list(doc.paragraphs[start_idx:]):
        remove_paragraph(paragraph)

    for style_name, text in SECTION_22:
        para = doc.add_paragraph(style=style_name)
        para.add_run(text)

    doc.save(str(TARGET))
    print(str(TARGET))


if __name__ == "__main__":
    main()

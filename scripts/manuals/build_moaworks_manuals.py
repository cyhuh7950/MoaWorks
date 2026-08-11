from __future__ import annotations

import re
from pathlib import Path
from PIL import Image

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
MANUAL_DIR = ROOT / "docs" / "manuals"
OUTPUT_DIR = MANUAL_DIR / "output"
SOURCES = [
    MANUAL_DIR / "moaworks-end-user-manual-v2.0.md",
    MANUAL_DIR / "moaworks-admin-operator-manual-v2.0.md",
    MANUAL_DIR / "moaworks-install-deploy-manual-v2.0.md",
    MANUAL_DIR / "moaworks-incident-backup-recovery-manual-v2.0.md",
]
NAVY, TEAL, GRAY, BORDER = "10233F", "087F75", "5F6B7A", "C9D5DF"


def font(run, size=None, color=None, bold=None, italic=None, name="Malgun Gothic"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def border(cell, color=BORDER):
    props = cell._tc.get_or_add_tcPr()
    borders = props.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        props.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:color"), color)
        borders.append(node)


def shade(cell, fill):
    props = cell._tc.get_or_add_tcPr()
    node = OxmlElement("w:shd")
    node.set(qn("w:fill"), fill)
    props.append(node)


def repeat_header(row):
    props = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    props.append(node)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])
    font(run, 9, GRAY)


def configure(doc, title):
    normal = doc.styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(NAVY)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.18

    for name, size, color in [
        ("Title", 24, NAVY), ("Heading 1", 18, NAVY),
        ("Heading 2", 14, TEAL), ("Heading 3", 11.5, NAVY),
        ("List Bullet", 10.5, NAVY), ("List Number", 10.5, NAVY),
    ]:
        style = doc.styles[name]
        style.font.name = "Malgun Gothic"
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        if name.startswith("Heading") or name == "Title":
            style.font.bold = True
            style.paragraph_format.keep_with_next = True

    section = doc.sections[0]
    section.page_width, section.page_height = Cm(21), Cm(29.7)
    section.top_margin, section.bottom_margin = Cm(1.8), Cm(1.7)
    section.left_margin, section.right_margin = Cm(1.8), Cm(1.8)
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header.add_run(title)
    font(r, 8.5, GRAY)
    add_page_number(section.footer.paragraphs[0])


def add_cover(doc, title, subtitle):
    doc.add_paragraph().paragraph_format.space_after = Pt(70)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(p.add_run("MOAWORKS GROUPWARE"), 11, TEAL, True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(p.add_run(title), 24, NAVY, True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(p.add_run(subtitle), 12, GRAY)
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    repeat_header(table.rows[0])
    cell = table.cell(0, 0)
    shade(cell, "E8F5F2")
    border(cell, TEAL)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(10)
    font(p.add_run("검증된 절차와 실제 화면 예시를 포함한 배포본"), 10.5, TEAL, True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(p.add_run("버전 2.0 · 2026-08-11"), 10, GRAY)
    p.add_run().add_break(WD_BREAK.PAGE)


def add_contents(doc, lines):
    doc.add_heading("목차", level=1)
    for line in lines:
        match = re.match(r"^(#{2,3})\s+(.+)$", line)
        if not match or len(match.group(1)) > 2:
            continue
        p = doc.add_paragraph()
        if len(match.group(1)) == 2:
            p.paragraph_format.left_indent = Cm(0.5)
        r = p.add_run(match.group(2))
        font(r, 10, NAVY if len(match.group(1)) == 2 else GRAY,
             len(match.group(1)) == 2)
    doc.add_page_break()


def add_text(paragraph, text):
    cleaned = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    cleaned = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", cleaned)
    cleaned = cleaned.replace(chr(96), "")
    font(paragraph.add_run(cleaned))


def add_table(doc, rows):
    width = max(len(row) for row in rows)
    table = doc.add_table(rows=0, cols=width)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row_index, data in enumerate(rows):
        row = table.add_row()
        if row_index == 0:
            repeat_header(row)
        for col in range(width):
            cell = row.cells[col]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            border(cell)
            if row_index == 0:
                shade(cell, NAVY)
            p = cell.paragraphs[0]
            add_text(p, data[col].strip() if col < len(data) else "")
            for run in p.runs:
                font(run, 8.7, "FFFFFF" if row_index == 0 else NAVY,
                     row_index == 0)
    doc.add_paragraph()


def add_image(doc, path, alt):
    if not path.exists():
        raise FileNotFoundError(path)
    with Image.open(path) as image:
        width_px, height_px = image.size
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if height_px / width_px > 1.4:
        shape = p.add_run().add_picture(str(path), height=Cm(20.5))
    else:
        shape = p.add_run().add_picture(str(path), width=Cm(16.8))
    shape._inline.docPr.set("title", alt)
    shape._inline.docPr.set("descr", f"MoaWorks 화면 예시: {alt}")
    p.paragraph_format.keep_together = True


def add_code(doc, code_lines):
    table = doc.add_table(rows=1, cols=1)
    repeat_header(table.rows[0])
    cell = table.cell(0, 0)
    shade(cell, "F3F6F8")
    border(cell)
    p = cell.paragraphs[0]
    for number, line in enumerate(code_lines):
        if number:
            p.add_run().add_break()
        font(p.add_run(line), 8.5, NAVY, name="Consolas")
    doc.add_paragraph()


def render(doc, source):
    lines = source.read_text(encoding="utf-8").splitlines()
    index, code = 0, None
    fence = chr(96) * 3
    while index < len(lines):
        line = lines[index].rstrip()
        if line.startswith(fence):
            if code is None:
                code = []
            else:
                add_code(doc, code)
                code = None
            index += 1
            continue
        if code is not None:
            code.append(line)
            index += 1
            continue

        if line.strip() == "<!-- pagebreak -->":
            doc.add_page_break()
            index += 1
            continue

        image = re.fullmatch(r"!\[([^\]]+)\]\(([^)]+)\)", line.strip())
        if image:
            add_image(doc, (source.parent / image.group(2)).resolve(), image.group(1))
            index += 1
            continue

        if line.startswith("|") and "|" in line[1:]:
            raw_rows = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                raw_rows.append(lines[index].strip())
                index += 1
            rows = []
            for raw in raw_rows:
                cells = [item.strip() for item in raw.strip("|").split("|")]
                if all(re.fullmatch(r":?-{3,}:?", item) for item in cells):
                    continue
                rows.append(cells)
            add_table(doc, rows)
            continue

        heading = re.match(r"^(#{1,3})\s+(.+)$", line)
        if heading:
            level = len(heading.group(1))
            if level > 1:
                doc.add_heading(heading.group(2), level=level - 1)
            index += 1
            continue

        if line.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            add_text(p, line[2:])
            for run in p.runs:
                font(run, 10, GRAY, italic=True)
        elif re.match(r"^- \[[ xX]\]", line):
            p = doc.add_paragraph(style="List Bullet")
            add_text(p, line.replace("- [ ]", "☐", 1).replace("- [x]", "☑", 1).replace("- [X]", "☑", 1))
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_text(p, line[2:])
        elif re.match(r"^\d+\.\s+", line):
            p = doc.add_paragraph(style="List Number")
            add_text(p, re.sub(r"^\d+\.\s+", "", line))
        elif line.startswith("*") and line.endswith("*") and len(line) > 2:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            font(p.add_run(line.strip("*")), 8.7, GRAY, italic=True)
        elif line.strip():
            add_text(doc.add_paragraph(), line.strip())
        index += 1


def build(source):
    lines = source.read_text(encoding="utf-8").splitlines()
    title = next(line[2:] for line in lines if line.startswith("# "))
    subtitles = {
        "moaworks-end-user-manual-v2.0.md": "일반 사용자를 위한 업무 기능 안내",
        "moaworks-admin-operator-manual-v2.0.md": "관리자와 운영자를 위한 업무·정책 안내",
        "moaworks-install-deploy-manual-v2.0.md": "설치 담당자를 위한 재현 가능한 구축 절차",
        "moaworks-incident-backup-recovery-manual-v2.0.md": "운영자를 위한 장애·백업·복구 절차",
    }
    doc = Document()
    configure(doc, title)
    add_cover(doc, title, subtitles[source.name])
    add_contents(doc, lines)
    render(doc, source)
    doc.core_properties.title = title
    doc.core_properties.author = "MoaWorks"
    doc.core_properties.subject = subtitles[source.name]
    doc.core_properties.keywords = "MoaWorks, groupware, manual"
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    output = OUTPUT_DIR / source.with_suffix(".docx").name
    doc.save(output)
    return output


def main():
    outputs = [build(source) for source in SOURCES]
    print(f"PASS: built {len(outputs)} manuals")
    for output in outputs:
        print(output.relative_to(ROOT))


if __name__ == "__main__":
    main()
